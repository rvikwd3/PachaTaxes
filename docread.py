print "Initialized"

from docx import Document
import re

# =====================================
# MAIN
# =====================================
def main():

    # Document object for file
    document = Document('InputFiles/GST.docx')

    # Number of tables
    print 'Tables:\t{}'.format(document.tables.__len__())

   # Iterate over each table
    for table_index in range(document.tables.__len__()):

        # GSTNumber Cell Index
        print "\nTable {}\tGST Index:\t{}\n".format(table_index, getGSTIndex(document.tables[table_index]))
        # Table text
        print "\nTable {} Text:\n{}".format(table_index, getTableText(document.tables[table_index]))

    # Details
    print "\n Table {} Details:\n{}".format('0', getDetails(document.tables[0]))

    # Expenses
    print "\n Table {} Expenses:\n{}".format('0', regexExpenses(getTableText(document.tables[0])))


# =====================================
# FUNCTIONS
# =====================================

#   1.  getTableText()
#   2.  getGSTIndex()


# =====================================
# getTableText()
# Text within given table

def getTableText(table):
    table_text = ""
    old_text = ""

    #   Number of rows & Columns
    print 'Rows:\t{}'.format(table.rows.__len__())
    table_rows = table.rows.__len__()
    print 'Cols:\t{}\n'.format(table.columns.__len__())
    table_cols = table.columns.__len__()

    #   Text within table
    for i in range(table_rows):
            for j in range(table_cols):

            #   (i,j) goes out of cell range for some reason
            #   we handle the "out of index range" error
                try:
                    new_text = table.cell(i,j).text
                except IndexError:
                    print '\t\tIndex Error at ({},{})'.format(i,j)
                    pass

                #Check if text is repeated
                if(old_text == new_text):
                    print "Repeated:\t{}".format(new_text)
                else:
                    print '({},{}):\t{}'.format(i,j,table.cell(i,j).text)
                    table_text += new_text+" "
                    old_text = new_text

    return table_text

# =====================================
# getGSTIndex()
# Index of cell containing GSTNumber
# Only checks first cell(0,0) of table

def getGSTIndex(table):
    print "\ngetGSTIndex"

    #   Check if GST number is in first cell
    gst_cell = table.cell(0,0).text
    gst_index = next(((index, i) for (index,i) in enumerate(gst_cell.split()) if re.search(r'GSTIN', gst_cell.split()[index], flags=re.M|re.S) is not None), (-1,-1))[0]

    #   If GST number wasn't found then gst_index <- -1
    if gst_index != -1:
            print '\nGST Index:\t{}'.format(gst_index)
            print '\nGST Number:\t{}'.format(gst_cell.split()[gst_index + 1])
    else: 
            print 'GST Index not found in {}'.format(filename)

    return gst_index

# =====================================
# getDetails()
# Details from first table of invoice

def getDetails(table):
    table_text = getTableText(table)

    #Multicolumn details
    multicolumn = [] 
    for i in range(table.columns.__len__()):
        multicolumn.append((table.cell(1,i).text, table.cell(2,i).text))

    multicolumn.append(('Date', regexDate(table_text)))
    multicolumn.append(('Inv#', regexInvoiceNumber(table_text)))
    
    return multicolumn

# =====================================
# REGEX FUNCTIONS
# =====================================

def regexDate(text_input):
    #Regex search for Dated:
    date_exists = re.search(r'Dated:\d{2}-[a-zA-z]+-\d{4}', text_input, flags=re.M|re.S)

    if date_exists:
        #String slice to extract date (6 is position of : in Dated:)
        date = date_exists.group(0)[6:].strip()
    else:
        date = 'N/A'

    return date

def regexInvoiceNumber(text_input):
    inv_number_exists = re.search(r'Invoice No : \d{10,11}', text_input, flags=re.M|re.S)

    if inv_number_exists:
        #String slice for invoice number (12 is position of : in Invoice No :)
        inv_no = inv_number_exists.group(0)[12:].strip()
    else:
        inv_no = 'N/A'

    return inv_no

def regexExpenses(text_input):
    cgst = re.search(r'CGST', text_input, flags=re.S|re.M)
    sgst = re.search(r'SGST', text_input, flags=re.S|re.M)
    igst = re.search(r'IGST', text_input, flags=re.S|re.M)
    total = re.search(r'Total Amount', text_input, flags=re.S|re.M)

    print cgst.group[0].strip()
    print sgst.group[0]
    print igst.group[0]
    print total.group[0]

    return "Heya"

# =====================================
# Runtime 
# =====================================

# File to read
filename = 'InputFiles/GST.docx'

main()